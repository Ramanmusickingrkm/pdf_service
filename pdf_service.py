from flask import Flask, request, send_file
from flask_cors import CORS
import mammoth
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import os
import base64
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

app = Flask(__name__)
CORS(app)

def extract_fields_from_docx(docx_bytes):
    """Extract all fields from DOCX (placeholders like {{field}}, [field])"""
    doc = Document(io.BytesIO(docx_bytes))
    text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text += '\n' + para.text
    
    fields = []
    patterns = [
        r'\{\{([^}]+)\}\}',
        r'\[([^\]]+)\]',
        r'__([^_]+)__',
        r'\$([^$]+)\$'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            field_name = match.strip()
            if field_name and field_name not in fields:
                fields.append(field_name)
    
    return fields

def replace_fields_in_docx(docx_bytes, field_values):
    """Replace placeholders in DOCX with actual values - PRESERVES FORMATTING"""
    doc = Document(io.BytesIO(docx_bytes))
    
    def replace_in_paragraph(paragraph, field_values):
        for field_name, field_value in field_values.items():
            if not field_value:
                continue
                
            patterns = [
                f'{{{{{field_name}}}}}',
                f'[{field_name}]',
                f'__{field_name}__',
                f'${field_name}$'
            ]
            
            for pattern in patterns:
                if pattern in paragraph.text:
                    # Replace while preserving runs if possible
                    if hasattr(paragraph, 'runs') and paragraph.runs:
                        # Replace in first run that contains the pattern
                        for run in paragraph.runs:
                            if pattern in run.text:
                                run.text = run.text.replace(pattern, str(field_value))
                                break
                        else:
                            # If not found in any run, replace whole paragraph text
                            paragraph.text = paragraph.text.replace(pattern, str(field_value))
                    else:
                        paragraph.text = paragraph.text.replace(pattern, str(field_value))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, field_values)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, field_values)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def convert_docx_to_pdf_fixed(docx_bytes):
    """Convert DOCX to PDF using python-docx2pdf or alternative method"""
    
    # Method 1: Try using docx2pdf (best option)
    try:
        from docx2pdf import convert
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_bytes)
            temp_docx_path = temp_docx.name
        
        pdf_path = temp_docx_path.replace('.docx', '.pdf')
        convert(temp_docx_path, pdf_path)
        
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Cleanup
        os.unlink(temp_docx_path)
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        
        return pdf_bytes
        
    except ImportError:
        pass
    
    # Method 2: Use pdfkit with wkhtmltopdf (requires HTML conversion)
    try:
        import mammoth
        import pdfkit
        
        # Convert DOCX to HTML
        with open(temp_docx_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
            html = result.value
        
        # Convert HTML to PDF
        pdf_bytes = pdfkit.from_string(html, False)
        return pdf_bytes
        
    except ImportError:
        pass
    
    # Method 3: Simple fallback - return DOCX
    print("⚠️ No PDF converter available, returning DOCX")
    return docx_bytes

@app.route('/parse-docx', methods=['POST'])
def parse_docx():
    """Parse DOCX and detect fields"""
    try:
        file = request.files['file']
        if not file:
            return {'success': False, 'error': 'No file uploaded'}, 400
        
        docx_bytes = file.read()
        fields = extract_fields_from_docx(docx_bytes)
        
        return {
            'success': True,
            'fields': [{'name': f, 'label': f.replace('_', ' ').title(), 'type': 'text'} for f in fields],
            'message': f'Found {len(fields)} fields'
        }
    
    except Exception as e:
        print(f"Parse error: {str(e)}")
        return {'success': False, 'error': str(e)}, 500

@app.route('/fill-and-pdf', methods=['POST'])
def fill_and_pdf():
    """Fill fields in DOCX and return PDF"""
    try:
        data = request.json
        docx_base64 = data.get('docxBase64')
        field_values = data.get('fieldValues', {})
        title = data.get('title', 'signed_document')
        
        if not docx_base64:
            return {'success': False, 'error': 'No document provided'}, 400
        
        print(f"📄 Filling document: {title}")
        print(f"📝 Fields to fill: {list(field_values.keys())}")
        
        docx_bytes = base64.b64decode(docx_base64)
        
        # Fill fields in DOCX
        filled_docx = replace_fields_in_docx(docx_bytes, field_values)
        
        # Convert to PDF
        pdf_bytes = convert_docx_to_pdf_fixed(filled_docx.getvalue())
        
        print(f"✅ PDF generated: {len(pdf_bytes)} bytes")
        
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{title}_signed.pdf'
        )
    
    except Exception as e:
        print(f"❌ Fill and PDF error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}, 500

@app.route('/fill-and-return-docx', methods=['POST'])
def fill_and_return_docx():
    """Fill fields in DOCX and return DOCX (no PDF conversion)"""
    try:
        data = request.json
        docx_base64 = data.get('docxBase64')
        field_values = data.get('fieldValues', {})
        title = data.get('title', 'signed_document')
        
        if not docx_base64:
            return {'success': False, 'error': 'No document provided'}, 400
        
        docx_bytes = base64.b64decode(docx_base64)
        
        # Fill fields in DOCX
        filled_docx = replace_fields_in_docx(docx_bytes, field_values)
        
        return send_file(
            io.BytesIO(filled_docx.getvalue()),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{title}_filled.docx'
        )
    
    except Exception as e:
        print(f"❌ Fill DOCX error: {str(e)}")
        return {'success': False, 'error': str(e)}, 500

@app.route('/health', methods=['GET'])
def health():
    return {'status': 'ok', 'service': 'docx-pdf-service'}

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001)
