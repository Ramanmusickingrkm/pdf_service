from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
import io
import re
import base64
import logging
from datetime import datetime

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def normalize_field_name(field_name):
    """Convert field name to standard format for matching"""
    # Remove brackets, underscores, convert to lowercase
    normalized = re.sub(r'[\[\]\{\}\$]', '', field_name)
    normalized = normalized.replace('_', ' ')
    normalized = normalized.lower().strip()
    return normalized

def get_field_mapping(docx_bytes):
    """Detect all fields in DOCX and create mapping to database field names"""
    doc = Document(io.BytesIO(docx_bytes))
    text = '\n'.join([para.text for para in doc.paragraphs])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text += '\n' + para.text
    
    # Find all placeholders
    fields = []
    patterns = [
        r'\{\{([^}]+)\}\}',
        r'\[([^\]]+)\]',
        r'__([^_]+)__',
        r'\$([^$]+)\$'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            field_name = match.strip()
            if field_name and field_name not in fields:
                fields.append(field_name)
    
    logger.info(f'🔍 Detected fields in template: {fields}')
    
    # Create mapping from template field to normalized form
    mapping = {}
    for field in fields:
        normalized = normalize_field_name(field)
        mapping[normalized] = field
        logger.info(f'  📌 "{field}" -> normalized: "{normalized}"')
    
    return mapping, fields

def replace_fields_in_docx(docx_bytes, field_values):
    """Replace placeholders in DOCX while preserving formatting"""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        # Get field mapping from template
        field_mapping, detected_fields = get_field_mapping(docx_bytes)
        
        logger.info(f'📝 Received field values: {list(field_values.keys())}')
        logger.info(f'📝 Detected template fields: {detected_fields}')
        
        # Create replacement dictionary
        replacements = {}
        
        for template_field in detected_fields:
            normalized_template = normalize_field_name(template_field)
            replaced = False
            
            # Try direct match
            if template_field in field_values:
                replacements[template_field] = str(field_values[template_field])
                logger.info(f'  ✓ Direct match: {template_field} -> {replacements[template_field][:30]}')
                replaced = True
            # Try normalized match (e.g., "organisation_name" matches "Organization Name")
            elif normalized_template in field_values:
                replacements[template_field] = str(field_values[normalized_template])
                logger.info(f'  ✓ Normalized match: {template_field} (normalized: {normalized_template}) -> {replacements[template_field][:30]}')
                replaced = True
            else:
                # Try case-insensitive match with spaces/underscores variations
                for key in field_values.keys():
                    # Create variations of database key
                    key_variations = [
                        key,  # organisation_name
                        key.replace('_', ' '),  # organisation name
                        key.replace('_', ' ').title(),  # Organisation Name
                        key.replace('_', ' ').capitalize(),  # Organisation name
                        key.upper(),  # ORGANISATION_NAME
                        key.replace('_', ' ').upper(),  # ORGANISATION NAME
                        key.split('_')[-1] if '_' in key else key,  # name
                    ]
                    
                    # Also try with 'z' vs 's' (organisation vs organization)
                    if 'organisation' in key:
                        key_variations.append(key.replace('organisation', 'organization'))
                        key_variations.append(key.replace('organisation', 'organization').replace('_', ' '))
                        key_variations.append(key.replace('organisation', 'organization').replace('_', ' ').title())
                    
                    for variation in key_variations:
                        if variation.lower() == template_field.lower():
                            replacements[template_field] = str(field_values[key])
                            logger.info(f'  ✓ Variation match: {template_field} -> {key} -> {replacements[template_field][:30]}')
                            replaced = True
                            break
                    if replaced:
                        break
            
            if not replaced:
                logger.warning(f'  ⚠️ No match found for field: {template_field}')
        
        if not replacements:
            logger.warning('⚠️ No field replacements found!')
            logger.warning(f'   Available field values: {list(field_values.keys())}')
            logger.warning(f'   Detected template fields: {detected_fields}')
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if not original_text:
                continue
                
            new_text = original_text
            for template_field, replacement_value in replacements.items():
                if not replacement_value:
                    continue
                
                # Try all pattern formats
                patterns_to_replace = [
                    f'{{{{{template_field}}}}}',
                    f'[{template_field}]',
                    f'__{template_field}__',
                    f'${template_field}$'
                ]
                
                for pattern in patterns_to_replace:
                    if pattern in new_text:
                        new_text = new_text.replace(pattern, replacement_value)
                        logger.info(f'  ✓ Replaced {pattern}')
            
            if new_text != original_text:
                # Try to preserve formatting by replacing in runs
                if paragraph.runs:
                    full_text = ''.join([run.text for run in paragraph.runs])
                    for template_field, replacement_value in replacements.items():
                        if not replacement_value:
                            continue
                        patterns_to_replace = [
                            f'{{{{{template_field}}}}}',
                            f'[{template_field}]',
                            f'__{template_field}__',
                            f'${template_field}$'
                        ]
                        for pattern in patterns_to_replace:
                            if pattern in full_text:
                                full_text = full_text.replace(pattern, replacement_value)
                    if full_text != original_text:
                        paragraph.clear()
                        paragraph.add_run(full_text)
                else:
                    paragraph.text = new_text
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if not original_text:
                            continue
                        new_text = original_text
                        for template_field, replacement_value in replacements.items():
                            if not replacement_value:
                                continue
                            patterns_to_replace = [
                                f'{{{{{template_field}}}}}',
                                f'[{template_field}]',
                                f'__{template_field}__',
                                f'${template_field}$'
                            ]
                            for pattern in patterns_to_replace:
                                if pattern in new_text:
                                    new_text = new_text.replace(pattern, replacement_value)
                        if new_text != original_text:
                            if paragraph.runs:
                                full_text = ''.join([run.text for run in paragraph.runs])
                                for template_field, replacement_value in replacements.items():
                                    if not replacement_value:
                                        continue
                                    patterns_to_replace = [
                                        f'{{{{{template_field}}}}}',
                                        f'[{template_field}]',
                                        f'__{template_field}__',
                                        f'${template_field}$'
                                    ]
                                    for pattern in patterns_to_replace:
                                        if pattern in full_text:
                                            full_text = full_text.replace(pattern, replacement_value)
                                if full_text != original_text:
                                    paragraph.clear()
                                    paragraph.add_run(full_text)
                            else:
                                paragraph.text = new_text
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        logger.info('✅ Fields replaced successfully')
        return output
        
    except Exception as e:
        logger.error(f'❌ Error replacing fields: {str(e)}')
        import traceback
        traceback.print_exc()
        return None

@app.route('/health', methods=['GET'])
def health():
    return {
        'status': 'ok',
        'service': 'DocSign DOCX Filler Service',
        'timestamp': datetime.now().isoformat()
    }

@app.route('/fill-and-pdf', methods=['POST'])
def fill_and_pdf():
    """Fill fields in DOCX and return filled DOCX"""
    try:
        data = request.json
        if not data:
            return {'success': False, 'error': 'No JSON data provided'}, 400
        
        docx_base64 = data.get('docxBase64')
        field_values = data.get('fieldValues', {})
        title = data.get('title', 'signed_document')
        
        logger.info(f'📊 Title: {title}')
        logger.info(f'📋 Field values received: {list(field_values.keys())}')
        
        if not docx_base64:
            return {'success': False, 'error': 'No document provided'}, 400
        
        try:
            docx_bytes = base64.b64decode(docx_base64)
            logger.info(f'✅ Decoded DOCX: {len(docx_bytes)} bytes')
        except Exception as e:
            return {'success': False, 'error': f'Failed to decode DOCX: {str(e)}'}, 400
        
        filled_docx_io = replace_fields_in_docx(docx_bytes, field_values)
        
        if not filled_docx_io:
            return {'success': False, 'error': 'Failed to replace fields'}, 500
        
        logger.info(f'✅ Returning filled DOCX: {len(filled_docx_io.getvalue())} bytes')
        
        return send_file(
            io.BytesIO(filled_docx_io.getvalue()),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{title}_filled.docx'
        )
    
    except Exception as e:
        logger.error(f'❌ Error: {str(e)}')
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}, 500

@app.route('/fill-and-return-docx', methods=['POST'])
def fill_and_return_docx():
    return fill_and_pdf()

if __name__ == '__main__':
    logger.info('🚀 Starting DocSign DOCX Filler Service')
    app.run(host='0.0.0.0', port=5001, debug=False)
